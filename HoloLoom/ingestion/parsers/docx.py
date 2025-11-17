"""
DOCX Parser - Extract Text from Word Documents
===============================================
Parse Microsoft Word .docx files.
"""

import logging
from pathlib import Path

logger = logging.getLogger(__name__)

# Try import with graceful fallback
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    logger.warning("python-docx not installed. DOCX parsing unavailable.")
    DOCX_AVAILABLE = False
    Document = None


async def parse_docx(file_path: str) -> str:
    """
    Extract text from DOCX file.

    Args:
        file_path: Path to DOCX file

    Returns:
        Extracted text

    Raises:
        RuntimeError: If DOCX parsing unavailable
    """
    if not DOCX_AVAILABLE:
        raise RuntimeError("DOCX parsing requires python-docx: pip install python-docx")

    path = Path(file_path)
    text_parts = []

    try:
        # Open document
        doc = Document(str(path))

        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    text_parts.append(row_text)

        # Combine
        full_text = "\n\n".join(text_parts)

        logger.info(f"Extracted {len(full_text)} characters from DOCX")

        return full_text

    except Exception as e:
        logger.error(f"Failed to parse DOCX {file_path}: {e}")
        raise


# ============================================================================
# Example Usage
# ============================================================================

if __name__ == "__main__":
    import asyncio

    async def demo():
        print("="*80)
        print("DOCX Parser Demo")
        print("="*80 + "\n")

        if not DOCX_AVAILABLE:
            print("python-docx not available")
            print("Install with: pip install python-docx")
            return

        print("DOCX parser is ready")
        print("Usage: text = await parse_docx('document.docx')")
        print("\n✓ Demo complete!")

    asyncio.run(demo())
